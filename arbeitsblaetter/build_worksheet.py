"""Datengesteuerter Arbeitsblatt-Generator für PyQuest im JJWS-Design.

Liest die Kapitel-Inhalte der App (public/content/) und erzeugt daraus
automatisch ein Information- & Aufgabenblatt als DOCX (und per LibreOffice
als PDF). So entsteht zu jedem in der App gepflegten Kapitel automatisch das
passende Arbeitsblatt - ohne pro Kapitel ein eigenes Skript.

Benutzung:
    python build_worksheet.py 02-variablen
    python build_worksheet.py --all
"""

import argparse
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

JJWS_SCRIPTS = (
    Path.home() / "AppData" / "Roaming" / "Claude" / "local-agent-mode-sessions"
    / "skills-plugin" / "e5c72aed-38e1-4d3f-b35e-9f62d76a831d"
    / "6dbbba93-6992-4a28-84a3-69f4f6ad7769" / "skills" / "jjws-design" / "scripts"
)
sys.path.insert(0, str(JJWS_SCRIPTS))
import jjws_brand as jj

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
CONTENT = ROOT / "public" / "content"
OUTDIR = Path(__file__).resolve().parent
PUBLIC_WORKSHEETS = ROOT / "public" / "worksheets"
SOFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"

NAVY = RGBColor.from_string(jj.NAVY)
BLUE = RGBColor.from_string(jj.BLUE)
GREY = RGBColor.from_string(jj.GREY)
CODE_FONT = "Consolas"
BASE_FONT = "Arial"  # Open Sans ggf. nicht installiert -> sicherer Fallback
CODE_FILL = "F2F5F7"


# ---------------------------------------------------------------- Low-level DOCX

def set_run_font(run, size=11, color=None, bold=False, font=None, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = font or BASE_FONT
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font or BASE_FONT)


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:keepNext"))
    return paragraph


def no_split_table(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))


def shade_paragraph(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


# ---------------------------------------------------------------- Markdown-ish

_BOLD = re.compile(r"(\*\*.+?\*\*)")
_CODE = re.compile(r"(`[^`]+?`)")


def add_rich_runs(paragraph, text, size=11, base_color=None):
    """Fügt Text mit **fett**, `code` und verschachteltem **`code`** als
    passend formatierte Runs ein."""
    for bold_part in _BOLD.split(text):
        if not bold_part:
            continue
        bold = bold_part.startswith("**") and bold_part.endswith("**")
        inner = bold_part[2:-2] if bold else bold_part
        # innerhalb (auch innerhalb von Fett) Inline-Code erkennen
        for piece in _CODE.split(inner):
            if not piece:
                continue
            if piece.startswith("`") and piece.endswith("`"):
                r = paragraph.add_run(piece[1:-1])
                set_run_font(r, size=size - 0.5, font=CODE_FONT, bold=bold,
                             color=NAVY if bold else base_color)
            else:
                r = paragraph.add_run(piece)
                set_run_font(r, size=size, bold=bold,
                             color=NAVY if bold else base_color)


def add_code_block(doc_or_cell, code, label=None, glue_next=False):
    if label:
        p = _add_paragraph(doc_or_cell)
        r = p.add_run(label)
        set_run_font(r, size=10, color=GREY, italic=True)
        p.paragraph_format.space_after = Pt(2)
        keep_with_next(p)
    p = _add_paragraph(doc_or_cell)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, CODE_FILL)
    for i, line in enumerate(code.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line if line else " ")
        set_run_font(r, size=10.5, font=CODE_FONT, color=NAVY)
    if glue_next:
        keep_with_next(p)
    return p


def _add_paragraph(doc_or_cell):
    return doc_or_cell.add_paragraph()


def render_markdown_block(doc, text, glue_next=True):
    """Rendert einen Erklär-/Infotext: Absätze, **fett**, `code`, ```codeblöcke```."""
    # Fenced Code-Bloecke herausziehen
    segments = re.split(r"```(?:python)?\n?(.*?)```", text, flags=re.DOTALL)
    for si, seg in enumerate(segments):
        if si % 2 == 1:  # Codeblock
            add_code_block(doc, seg.rstrip("\n"))
            continue
        for para in re.split(r"\n{2,}", seg):
            para = para.strip("\n")
            if not para.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            # einfache Zeilenumbrueche innerhalb eines Absatzes erhalten
            for li, line in enumerate(para.split("\n")):
                if li > 0:
                    p.add_run().add_break()
                add_rich_runs(p, line)
            keep_with_next(p)


# ---------------------------------------------------------------- Bausteine

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 12.5, color=NAVY, bold=True)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), jj.BLUE)
        pBdr.append(bottom); pPr.append(pBdr)
    keep_with_next(p)
    return p


def new_task_box(doc, number, question_text):
    """Erzeugt eine Aufgaben-Box (Markenblau-Rahmen, heller Hintergrund) und
    rendert die Aufgabenstellung richtig: die Leitzeile als „Aufgabe N: …“ in
    Navy fett, Code-Blöcke sauber als Code, weiterer Text normal.
    Gibt die Zelle zum Weiterfüllen (Antwortzeilen/Optionen) zurück."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    jj._set_table_borders(table, jj.BLUE, sz=8)
    jj._set_cell_shading(cell, jj.BOX_FILL)
    jj._set_cell_margins(cell)
    no_split_table(table)

    segments = re.split(r"```(?:python)?\n?(.*?)```", question_text, flags=re.DOTALL)
    first_line = True
    for si, seg in enumerate(segments):
        if si % 2 == 1:  # Code-Block
            add_code_block(cell, seg.rstrip("\n"))
            continue
        for para in [x for x in re.split(r"\n{2,}", seg) if x.strip()]:
            p = cell.paragraphs[0] if first_line and not cell.paragraphs[0].runs else cell.add_paragraph()
            if first_line:
                lbl = p.add_run(f"Aufgabe {number}: ")
                set_run_font(lbl, size=11, bold=True, color=NAVY)
                first_line = False
                # Leitzeile ebenfalls fett (wie eine Aufgaben-Ueberschrift)
                _rich_bold(p, para.replace("\n", " "))
            else:
                for li, line in enumerate(para.split("\n")):
                    if li > 0:
                        p.add_run().add_break()
                    add_rich_runs(p, line, size=11)
    if first_line:  # falls die Aufgabe leer/nur Code war -> trotzdem Label
        p = cell.paragraphs[0]
        lbl = p.add_run(f"Aufgabe {number}:")
        set_run_font(lbl, size=11, bold=True, color=NAVY)
    return cell


def _rich_bold(paragraph, text, size=11):
    """Leitzeile einer Aufgabe: normal fett-navy, `code` als fette Monospace.
    Die ganze Zeile ist bereits fett, daher werden **…**-Marker nur entfernt."""
    text = text.replace("**", "")
    for piece in _CODE.split(text):
        if not piece:
            continue
        if piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            set_run_font(r, size=size - 0.5, font=CODE_FONT, bold=True, color=NAVY)
        else:
            r = paragraph.add_run(piece)
            set_run_font(r, size=size, bold=True, color=NAVY)


def add_choice_task(doc, number, question, choices):
    cell = new_task_box(doc, number, question)
    letters = "abcdefgh"
    cell.add_paragraph()
    for i, c in enumerate(choices):
        p = cell.add_paragraph()
        set_run_font(p.add_run(f"☐  {letters[i]})  "), size=11)
        add_rich_runs(p, str(c), size=11)


def add_code_task(doc, number, task_text, starter=None, answer_lines=3, template=None):
    cell = new_task_box(doc, number, task_text)
    if template:
        p = cell.add_paragraph()
        r = p.add_run(re.sub(r"`", "", template))
        set_run_font(r, size=11, font=CODE_FONT, color=NAVY)
    if starter and starter.strip():
        p = cell.add_paragraph()
        for i, line in enumerate(starter.rstrip("\n").split("\n")):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(line if line else " ")
            set_run_font(r, size=10, font=CODE_FONT, color=GREY, italic=True)
    p = cell.add_paragraph()
    set_run_font(p.add_run("Dein Code:"), size=9.5, color=GREY, italic=True)
    for _ in range(answer_lines):
        cell.add_paragraph()


# ---------------------------------------------------------------- Kopf/Fuß

def add_horizontal_header(doc, fach_value="IT", logo_height_in=0.5, line_len=13):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    name_cell, fach_cell, logo_cell = table.rows[0].cells
    name_cell.width = Inches(3.6); fach_cell.width = Inches(1.4); logo_cell.width = Inches(1.5)
    p = name_cell.paragraphs[0]
    set_run_font(p.add_run("Name: "), size=10, bold=True)
    set_run_font(p.add_run("_" * line_len), size=10)
    p = fach_cell.paragraphs[0]
    set_run_font(p.add_run("Fach: "), size=10, bold=True)
    set_run_font(p.add_run(fach_value), size=10)
    p = logo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(jj.logo_path("bildmarke", "bunt"), height=Inches(logo_height_in))
    jj._clear_table_borders(table)


def _page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=GREY)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def add_footer(doc, left="JJWS", center="Riegert"):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc, cc, rc = table.rows[0].cells
    lc.width = Inches(2.17); cc.width = Inches(2.16); rc.width = Inches(2.17)
    set_run_font(lc.paragraphs[0].add_run(left), size=9, color=GREY, bold=True)
    cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cc.paragraphs[0].add_run(center), size=9, color=GREY)
    rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(rc.paragraphs[0].add_run("Seite "), size=9, color=GREY)
    _page_field(rc.paragraphs[0])
    jj._clear_table_borders(table)


# ---------------------------------------------------------------- Content laden

def load_curriculum():
    return json.loads((CONTENT / "curriculum.json").read_text(encoding="utf-8"))


def load_chapter(chapter_id):
    ch = json.loads((CONTENT / "chapters" / chapter_id / "chapter.json").read_text(encoding="utf-8"))
    lessons = []
    for lid in ch["lessons"]:
        lp = CONTENT / "chapters" / chapter_id / "lessons" / f"{lid}.json"
        lessons.append(json.loads(lp.read_text(encoding="utf-8")))
    return ch, lessons


def chapter_number(curriculum, chapter_id):
    return curriculum["chapters"].index(chapter_id) + 1


# ---------------------------------------------------------------- Dokument bauen

def build_docx(chapter_id):
    curriculum = load_curriculum()
    ch, lessons = load_chapter(chapter_id)
    num = chapter_number(curriculum, chapter_id)

    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.7)

    jj.set_docx_base_font(doc, primary="Arial")
    add_horizontal_header(doc, fach_value="IT")
    add_footer(doc)

    # Titelblock
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("PyQuest"), size=13, color=BLUE, bold=True); keep_with_next(p)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("Information- & Aufgabenblatt"), size=22, color=NAVY, bold=True); keep_with_next(p)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(f"Kapitel {num}: {ch['title']}"), size=14, color=BLUE, bold=True); keep_with_next(p)

    if ch.get("description"):
        pp = doc.add_paragraph(); pp.paragraph_format.space_after = Pt(8)
        add_rich_runs(pp, ch["description"])

    task_no = 0
    for lesson in lessons:
        add_heading(doc, lesson["title"], level=1)
        for step in lesson["steps"]:
            t = step["type"]
            if t == "explain":
                render_markdown_block(doc, step["text"])
            elif t == "bauplan":
                # Beschriftete Zerlegung einer Codezeile (siehe lesson-view.js).
                # Auf Papier: Codezeile als Block, darunter die nummerierte
                # Legende - dieselbe Reihenfolge wie in der App.
                if step.get("titel"):
                    render_markdown_block(doc, f"**{step['titel']}**")
                if step.get("text"):
                    render_markdown_block(doc, step["text"])
                zeile = "".join(p["text"] for p in step.get("teile", []))
                if step.get("folgezeile"):
                    zeile += "\n    " + step["folgezeile"]
                add_code_block(doc, zeile, label="Bauplan")
                nr = 0
                for p in step.get("teile", []):
                    if not p.get("name"):
                        continue
                    nr += 1
                    render_markdown_block(doc, f"{nr}. **{p['name']}** – {p.get('erklaerung', '')}")
                if step.get("folgeName"):
                    render_markdown_block(doc, f"{nr + 1}. **{step['folgeName']}**")
            elif t == "example":
                if step.get("text"):
                    render_markdown_block(doc, step["text"])
                add_code_block(doc, step["code"], label="Beispiel")
            elif t == "quiz":
                task_no += 1
                add_choice_task(doc, task_no, step["question"], step["choices"])
            elif t == "fill":
                task_no += 1
                add_code_task(doc, task_no, step.get("text", "Ergänze den Code:"),
                              template=step.get("template"), answer_lines=1)
            elif t == "code":
                task_no += 1
                add_code_task(doc, task_no, step["task"], starter=step.get("starterCode"),
                              answer_lines=3)

    out = OUTDIR / f"Kapitel_{num:02d}_{safe(ch['title'])}_Information_Aufgabenblatt.docx"
    doc.save(out)
    return out


def safe(name):
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")


def to_pdf(docx_path):
    subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf",
                    "--outdir", str(docx_path.parent), str(docx_path)],
                   check=True, capture_output=True)
    return docx_path.with_suffix(".pdf")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("chapter", nargs="?", help="Kapitel-ID, z. B. 02-variablen")
    ap.add_argument("--all", action="store_true", help="alle Kapitel bauen")
    ap.add_argument("--no-pdf", action="store_true", help="nur DOCX, kein PDF")
    args = ap.parse_args()

    curriculum = load_curriculum()
    targets = curriculum["chapters"] if args.all else [args.chapter]
    if not targets or targets == [None]:
        ap.error("Bitte eine Kapitel-ID angeben oder --all benutzen.")

    for cid in targets:
        docx = build_docx(cid)
        print("DOCX:", docx.name)
        if not args.no_pdf:
            pdf = to_pdf(docx)
            print("PDF :", pdf.name)
            PUBLIC_WORKSHEETS.mkdir(parents=True, exist_ok=True)
            published = PUBLIC_WORKSHEETS / f"{cid}.pdf"
            shutil.copyfile(pdf, published)
            print("     ->", published.relative_to(ROOT))


if __name__ == "__main__":
    main()
