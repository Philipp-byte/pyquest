"""Prueft, ob die erzeugten Arbeitsblaetter zum Kursinhalt passen.

Die Arbeitsblaetter (.docx, .pdf) sind versioniert und werden aus den
Lektionen erzeugt - aber nicht automatisch. Beim Gesamtcheck fiel auf,
dass nach Inhaltsaenderungen in vier Kapiteln die Blaetter nicht neu
erzeugt worden waren; auf GitHub Pages lagen veraltete Blaetter ohne die
neuen Aufgaben, und nichts hatte es gemeldet.

Zwei Dinge werden je Kapitel geprueft:

1. Die Zahl der Aufgaben im Blatt stimmt mit der Zahl der Aufgaben in den
   Lektionen ueberein (quiz + fill + code). Kommt eine Aufgabe dazu und das
   Blatt bleibt alt, faellt es hier auf.
2. Jeder Bauplan-Titel steht im Blatt. Faellt ein Schritt-Typ durch die
   Renderer, faellt es hier auf.

Gelesen wird die .docx-Quelle inklusive Tabellenzellen - dort stehen die
Aufgaben; wer nur Document.paragraphs liest, sieht sie nicht.

Aufruf:  python pruefe_arbeitsblaetter.py
"""

import json
import re
import sys
from pathlib import Path

from docx import Document

HIER = Path(__file__).resolve().parent
CONTENT = HIER.parent / "public" / "content"
AUFGABEN_TYPEN = {"quiz", "fill", "code"}


def docx_text(pfad):
    d = Document(pfad)
    teile = [p.text for p in d.paragraphs]
    for tabelle in d.tables:
        for zeile in tabelle.rows:
            for zelle in zeile.cells:
                teile.append(zelle.text)
    return "\n".join(teile)


def safe(name):
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")


def main():
    curriculum = json.loads((CONTENT / "curriculum.json").read_text(encoding="utf-8"))
    fehler = []
    geprueft = 0

    for nr, cid in enumerate(curriculum["chapters"], start=1):
        chapter = json.loads((CONTENT / "chapters" / cid / "chapter.json").read_text(encoding="utf-8"))
        erwartet = 0
        bauplaene = []
        for lid in chapter["lessons"]:
            lesson = json.loads((CONTENT / "chapters" / cid / "lessons" / f"{lid}.json").read_text(encoding="utf-8"))
            for step in lesson["steps"]:
                if step.get("type") in AUFGABEN_TYPEN:
                    erwartet += 1
                if step.get("type") == "bauplan" and step.get("titel"):
                    bauplaene.append(step["titel"])

        docx = HIER / f"Kapitel_{nr:02d}_{safe(chapter['title'])}_Information_Aufgabenblatt.docx"
        pdf = HIER.parent / "public" / "worksheets" / f"{cid}.pdf"
        geprueft += 1

        if not docx.exists():
            fehler.append(f"Kapitel {nr}: {docx.name} fehlt")
            continue
        if not pdf.exists():
            fehler.append(f"Kapitel {nr}: public/worksheets/{cid}.pdf fehlt")

        text = docx_text(docx)
        gefunden = len(set(re.findall(r"Aufgabe (\d+)", text)))
        if gefunden != erwartet:
            fehler.append(f"Kapitel {nr}: Blatt hat {gefunden} Aufgaben, Lektionen haben {erwartet} "
                          f"- Blatt neu erzeugen (python build_worksheet.py {cid})")
        for titel in bauplaene:
            if titel not in text:
                fehler.append(f"Kapitel {nr}: Bauplan '{titel}' fehlt im Blatt")

    for f in fehler:
        print("FEHLER:", f)
    print(f"--- {geprueft} Arbeitsblaetter geprueft, {len(fehler)} Fehler ---")
    return 1 if fehler else 0


if __name__ == "__main__":
    sys.exit(main())
